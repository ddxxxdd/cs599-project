"""Generate the CS599 Word report for Aliyun OSS Support RAG."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from aliyun_oss_rag.config import DOCS_DIR  # noqa: E402

OUT = DOCS_DIR / "CS599_大作业报告.docx"
BLUE = RGBColor(31, 111, 235)
DARK_BLUE = RGBColor(19, 61, 118)
LIGHT_FILL = "F2F6FB"


def set_run_font(run, western: str = "Calibri", east_asia: str = "宋体", size: int | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = western
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, western: str = "Calibri", east_asia: str = "宋体") -> None:
    style.font.name = western
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    set_style_font(normal, east_asia="宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        st = doc.styles[name]
        set_style_font(st, east_asia="黑体")
        st.font.size = Pt(size)
        st.font.color.rgb = color


def add_para(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("企业级应用软件设计与开发")
    set_run_font(r, east_asia="黑体", size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Aliyun OSS Support RAG：阿里云对象存储技术支持问答智能体")
    set_run_font(r, east_asia="黑体", size=17, bold=True)
    add_table(
        doc,
        ["字段", "内容"],
        [
            ["课程名称", "企业级应用软件设计与开发"],
            ["项目名称", "Aliyun OSS Support RAG"],
            ["方向", "方向一：Agentic AI 原生开发"],
            ["学号", "待填写"],
            ["姓名", "待填写"],
            ["专业", "计算机技术 / 软件工程（待替换）"],
            ["指导教师", "戚欣"],
            ["提交日期", "2026 年 6 月 22 日"],
        ],
    )
    add_para(doc, "说明：封面中的姓名、学号、专业为占位信息，正式提交前必须替换为真实信息。")
    doc.add_page_break()


def eval_summary() -> dict[str, object]:
    path = PROJECT_ROOT / "data" / "eval" / "results.json"
    if not path.exists():
        return {"total": "待运行", "intent_accuracy": "待运行", "must_include_accuracy": "待运行", "citation_rate": "待运行"}
    return json.loads(path.read_text(encoding="utf-8"))["summary"]


def kb_summary() -> dict[str, object]:
    return json.loads((PROJECT_ROOT / "data" / "processed" / "retrieval_index.json").read_text(encoding="utf-8"))


def main() -> None:
    doc = Document()
    style_doc(doc)
    add_cover(doc)

    add_para(doc, "一、选题背景与设计思想", style="Heading 1")
    add_para(doc, "企业应用接入阿里云 OSS 时常遇到权限策略、临时凭证、API 参数、签名错误、跨域和生命周期成本等问题。通用大模型可能给出宽泛答案，但技术支持场景需要基于具体官方文档、错误码和操作边界回答。")
    add_para(doc, "本项目从零构建阿里云 OSS 技术支持 RAG Agent，通过本地知识库检索、工具调用、状态编排和流式 Web 工作台，提高回答的可追溯性和工程可用性。")

    add_para(doc, "二、系统架构", style="Heading 1")
    add_table(doc, ["层级", "职责"], [["入口层", "React OSS 支持工作台 / FastAPI"], ["Agent 层", "classify -> retrieve -> tools -> synthesize"], ["工具层", "doc_lookup、api_lookup、permission_lookup、troubleshoot_lookup、cost_lookup"], ["数据层", "OSS 文档摘要、60 个 chunks、文档索引、向量索引、benchmark"]])
    add_para(doc, "FastAPI 同时托管前端构建产物和后端接口。LLM 层通过 .env 中的 BASE_URL、API_KEY、MODEL 三项连接 OpenAI 兼容 Chat Completions API；Embedding 层通过 EMBEDDING_BASE_URL、EMBEDDING_API_KEY、EMBEDDING_MODEL 连接 OpenAI 兼容 Embeddings API，未配置时回退到 BM25。")

    add_para(doc, "三、知识库构建", style="Heading 1")
    k = kb_summary()
    add_table(doc, ["指标", "结果"], [["产品", str(k["product"])], ["文档数", str(k["document_count"])], ["chunk 数", str(k["chunk_count"])], ["主题数", str(k["topic_count"])], ["来源", "、".join(k["sources"])]])
    add_para(doc, "知识库只包含阿里云 OSS 技术支持知识和官方来源 URL，不包含课程要求、项目说明或报告文本。")

    add_para(doc, "四、关键实现", style="Heading 1")
    add_table(doc, ["模块", "关键职责"], [["build_kb.py", "读取 aliyun_oss_docs.json，按概览、操作、权限、API、故障和成本章节生成 chunks，并在 Embedding 配置存在时生成 vector_index.json"], ["retrieval.py", "执行向量相似度 + BM25 混合检索，未配置 Embedding 时自动回退 BM25"], ["agent.py", "LangGraph 编排、意图识别、工具调用、引用片段和回答合成"], ["tools.py", "文档查询、API 查询、权限指南、故障排查和主题筛选"], ["frontend/src/main.tsx", "React 技术支持工作台、流式展示、历史列表、知识库依据和引用跳转"], ["evaluate.py", "OSS benchmark 与指标输出"]])

    add_para(doc, "五、测试与评估", style="Heading 1")
    s = eval_summary()
    add_table(doc, ["指标", "结果"], [["样例数", str(s["total"])], ["意图识别准确率", str(s["intent_accuracy"])], ["关键字段覆盖率", str(s["must_include_accuracy"])], ["引用覆盖率", str(s["citation_rate"])]])
    add_para(doc, "测试覆盖 AccessDenied、STS、Bucket Policy、API、签名错误、CORS、生命周期、加密和版本控制等 OSS 技术支持场景。")

    add_para(doc, "六、扩展方向", style="Heading 1")
    add_bullets(doc, ["接入阿里云 ECS、RAM、CDN、RDS 等更多产品文档。", "加入真实文档爬取、增量更新、向量索引增量刷新和重排序模型。", "扩展为权限 Agent、故障排查 Agent、API Agent 和成本优化 Agent 的多智能体系统。"])

    add_para(doc, "知识库资料来源", style="Heading 1")
    add_table(doc, ["来源", "链接"], [["阿里云 OSS 官方文档", "https://www.alibabacloud.com/help/zh/oss/"], ["阿里云 OpenAPI 文档", "https://api.alibabacloud.com/"]])

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    main()
